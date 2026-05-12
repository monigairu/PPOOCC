"""
Layer 1: 文書パーサー

委託会社から提出された Word/Excel ファイルを読み込み、
構造を保持したテキスト形式に変換する。

LLM は一切使用しない（決定論的処理）。
"""
from pathlib import Path

from openpyxl import load_workbook


def parse_file(file_path: str) -> str:
    """
    ファイルを読み込み、構造化テキストとして返す。

    対応形式:
      - .xlsx / .xls → Excel パーサー
      - .docx → Word パーサー

    Args:
        file_path: 読み込むファイルのパス

    Returns:
        構造化テキスト（mapper に渡す中間表現）
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"ファイルが見つかりません: {file_path}")

    suffix = path.suffix.lower()

    if suffix in (".xlsx", ".xls"):
        return _parse_excel(path)
    elif suffix == ".docx":
        return _parse_word(path)
    else:
        raise ValueError(
            f"未対応のファイル形式です: {suffix}\n"
            f"対応形式: .xlsx, .xls, .docx"
        )


def _parse_excel(path: Path) -> str:
    """
    Excel ファイルを構造化テキストに変換する。

    各シートの内容を以下の形式で出力:
      ## シート: シート名
      行番号 | セルA | セルB | セルC | ...

    結合セルは親セルの値を使い、空セルは "(空)" と表示する。
    """
    wb = load_workbook(str(path), read_only=True, data_only=True)
    output_lines: list[str] = []

    output_lines.append(f"# ファイル: {path.name}")
    output_lines.append(f"# 形式: Excel")
    output_lines.append("")

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        output_lines.append(f"## シート: {sheet_name}")
        output_lines.append("")

        for row_idx, row in enumerate(ws.iter_rows(), start=1):
            cells: list[str] = []

            for cell in row:
                # read_only モードでは MergedCell ではなく EmptyCell が返されるため
                # hasattr で value の有無を確認してから処理する
                if not hasattr(cell, "value") or cell.value is None:
                    cells.append("(空)")
                else:
                    cells.append(str(cell.value).strip())

            # 全セルが空または結合のみの行はスキップ
            if all(c in ("(空)", "(結合)") for c in cells):
                continue

            output_lines.append(f"行{row_idx} | {'  |  '.join(cells)}")

        output_lines.append("")

    wb.close()
    return "\n".join(output_lines)


def _parse_word(path: Path) -> str:
    """
    Word ファイルを構造化テキストに変換する。

    段落と表を区別して出力:
      - 段落: テキストをそのまま出力（見出しレベル付き）
      - 表: 行/列をパイプ区切りで出力

    python-docx が必要。
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx がインストールされていません。\n"
            "pip install python-docx を実行してください。"
        )

    doc = Document(str(path))
    output_lines: list[str] = []

    output_lines.append(f"# ファイル: {path.name}")
    output_lines.append(f"# 形式: Word")
    output_lines.append("")

    # 段落と表が文書内で混在するため、body の要素を順に処理
    table_index = 0
    tables = doc.tables

    for element in doc.element.body:
        tag = element.tag.split("}")[-1]  # 名前空間を除去

        if tag == "p":
            # 段落の処理
            paragraph = _find_paragraph_by_element(doc, element)
            if paragraph is not None and paragraph.text.strip():
                style_name = paragraph.style.name if paragraph.style else ""

                if "Heading" in style_name:
                    level = style_name.replace("Heading ", "").strip()
                    try:
                        level_num = int(level)
                    except ValueError:
                        level_num = 1
                    prefix = "#" * (level_num + 1)
                    output_lines.append(
                        f"{prefix} {paragraph.text.strip()}"
                    )
                else:
                    output_lines.append(paragraph.text.strip())

        elif tag == "tbl":
            # 表の処理
            if table_index < len(tables):
                table = tables[table_index]
                output_lines.append("")
                output_lines.append(f"### 表{table_index + 1}")

                for row_idx, row in enumerate(table.rows):
                    cells = [
                        cell.text.strip() if cell.text.strip() else "(空)"
                        for cell in row.cells
                    ]
                    output_lines.append(
                        f"行{row_idx + 1} | {'  |  '.join(cells)}"
                    )

                output_lines.append("")
                table_index += 1

    return "\n".join(output_lines)


def _find_paragraph_by_element(doc, element):
    """
    XML 要素から対応する Paragraph オブジェクトを探す。
    """
    for paragraph in doc.paragraphs:
        if paragraph._element is element:
            return paragraph
    return None
